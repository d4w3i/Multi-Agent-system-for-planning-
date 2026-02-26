"""
=============================================================================
TOOLS.PY - Function Tools for AI Agents
=============================================================================

This module provides function tools for AI agents, allowing them to:
1. Analyze source code and build call graphs
2. Navigate the file system safely
3. Read files of any format (code, documents, images)

ARCHITECTURE:

    Tools are organized in two layers:

    ┌─────────────────────────────────────────────────────────────┐
    │                    @function_tool Wrappers                  │
    │  (read_file, list_directory, find_code_files, etc.)         │
    └───────────────────────────┬─────────────────────────────────┘
                                │ calls
    ┌───────────────────────────▼─────────────────────────────────┐
    │                    _impl Functions                          │
    │  (_read_file_impl, _list_directory_impl)                    │
    │  Core logic, testable without decorator                     │
    └───────────────────────────┬─────────────────────────────────┘
                                │ uses
    ┌───────────────────────────▼─────────────────────────────────┐
    │                    Helper Functions                         │
    │  (_read_text_file, _read_pdf, _detect_file_type, etc.)      │
    └─────────────────────────────────────────────────────────────┘

AVAILABLE TOOLS:

    get_function_context    Analyze a function with its dependencies
    read_file               Read any file type (code, PDF, Excel, images)
    list_directory          List directory contents with filters
    find_code_files         Find files by programming language

SECURITY:
    - All tools block path traversal (..)
    - Optional sandboxing support with base_dir
    - Configurable file size limits (default 50MB)

USAGE:

    # As function tools (for OpenAI Agents SDK)
    from GenAI.tools import read_file, list_directory

    agent = Agent(tools=[read_file, list_directory])

    # Direct use of implementation functions
    from GenAI.tools import _read_file_impl

    content = _read_file_impl("main.py", verbose=False)

=============================================================================
"""

from typing import Optional
from pathlib import Path
from agents import function_tool
from colorama import Fore, Style, init

# Initialize colorama
init(autoreset=True)

# Error prefix constant for reliable error detection
_ERROR_PREFIX = "\u274c"  # ❌ emoji


def _is_error_result(result: str) -> bool:
    """Check if a result string indicates an error (starts with ❌)."""
    return result.startswith(_ERROR_PREFIX)


# =============================================================================
# SECTION 2: FILE SYSTEM TOOLS

# This section contains tools for file system navigation and reading
# files in various formats. These tools are designed to be secure and
# provide structured output useful for AI agents.

# Architecture:
# - _impl functions: core implementation (testable, without decorator)
# - @function_tool functions: wrappers for use with agents
# - Helper functions: _read_*, _detect_* for specific format handling
#
# =============================================================================

# -----------------------------------------------------------------------------
# CONSTANTS: Extension -> Language Mapping
# These constants define how to identify file types based on
# extension. They are used by _detect_file_type() to determine
# how to process each file.
CODE_EXTENSIONS: dict[str, str] = {
    # Python
    '.py': 'python',
    '.pyw': 'python',
    '.pyi': 'python-stub',
    '.pyx': 'cython',

    # Shell (common in Python project tooling)
    '.sh': 'bash',
    '.bash': 'bash',

    # Data / Configuration
    '.json': 'json',
    '.yaml': 'yaml',
    '.yml': 'yaml',
    '.toml': 'toml',
    '.ini': 'ini',
    '.cfg': 'ini',
    '.conf': 'conf',
    '.xml': 'xml',
    '.csv': 'csv',

    # Documentation
    '.md': 'markdown',
    '.rst': 'restructuredtext',
    '.txt': 'text',

    # SQL (common in Python backend projects)
    '.sql': 'sql',

    # Web (common in Django/Flask/FastAPI projects)
    '.html': 'html',
    '.htm': 'html',
    '.css': 'css',
    '.js': 'javascript',

    # Dockerfile (common in Python deployment)
    '.dockerfile': 'dockerfile',
    '.containerfile': 'dockerfile',
}

# -----------------------------------------------------------------------------
# CONSTANTS: Special Files (without extension or with specific name)
# Some important files have no extension or have specific names.
# This map allows identifying them correctly.
SPECIAL_FILES: dict[str, str] = {
    # DevOps / Container
    'Dockerfile': 'dockerfile',
    'Containerfile': 'dockerfile',
    'docker-compose.yml': 'yaml',
    'docker-compose.yaml': 'yaml',

    # Build systems
    'Makefile': 'makefile',
    'CMakeLists.txt': 'cmake',
    'meson.build': 'meson',
    'BUILD': 'bazel',
    'BUILD.bazel': 'bazel',
    'WORKSPACE': 'bazel',

    # Ruby ecosystem
    'Gemfile': 'ruby',
    'Rakefile': 'ruby',
    'Guardfile': 'ruby',
    'Vagrantfile': 'ruby',

    # CI/CD
    'Jenkinsfile': 'groovy',
    '.gitlab-ci.yml': 'yaml',
    '.travis.yml': 'yaml',

    # Platform specific
    'Procfile': 'procfile',
    'Brewfile': 'ruby',

    # Git
    '.gitignore': 'gitignore',
    '.gitattributes': 'gitattributes',
    '.gitmodules': 'gitmodules',

    # Docker
    '.dockerignore': 'dockerignore',

    # Environment
    '.env': 'dotenv',
    '.env.local': 'dotenv',
    '.env.development': 'dotenv',
    '.env.production': 'dotenv',
    '.env.example': 'dotenv',

    # Editor config
    '.editorconfig': 'editorconfig',
    '.prettierrc': 'json',
    '.eslintrc': 'json',
    '.stylelintrc': 'json',

    # Python ecosystem
    'requirements.txt': 'requirements',
    'constraints.txt': 'requirements',
    'Pipfile': 'toml',
    'Pipfile.lock': 'json',
    'pyproject.toml': 'toml',
    'setup.py': 'python',
    'setup.cfg': 'ini',
    'tox.ini': 'ini',
    'pytest.ini': 'ini',
    '.flake8': 'ini',
    'MANIFEST.in': 'manifest',

    # Node.js ecosystem
    'package.json': 'json',
    'package-lock.json': 'json',
    'tsconfig.json': 'jsonc',
    'jsconfig.json': 'jsonc',
    '.npmrc': 'ini',
    '.nvmrc': 'text',
    'yarn.lock': 'yaml',
    'pnpm-lock.yaml': 'yaml',

    # PHP
    'composer.json': 'json',
    'composer.lock': 'json',

    # Rust
    'Cargo.toml': 'toml',
    'Cargo.lock': 'toml',

    # Go
    'go.mod': 'gomod',
    'go.sum': 'gosum',
    'go.work': 'gowork',

    # License / Docs
    'LICENSE': 'text',
    'LICENSE.md': 'markdown',
    'LICENSE.txt': 'text',
    'README': 'text',
    'README.md': 'markdown',
    'README.txt': 'text',
    'CHANGELOG': 'text',
    'CHANGELOG.md': 'markdown',
    'CONTRIBUTING': 'text',
    'CONTRIBUTING.md': 'markdown',
    'AUTHORS': 'text',
    'CODEOWNERS': 'codeowners',
}

# -----------------------------------------------------------------------------
# CONSTANTS: Document Formats (binaries that require special parsing)
DOCUMENT_EXTENSIONS: dict[str, str] = {
    # PDF
    '.pdf': 'pdf',

    # Microsoft Office
    '.doc': 'word',      # Word legacy
    '.docx': 'word',     # Modern Word (Office Open XML)
    '.xls': 'excel',     # Excel legacy
    '.xlsx': 'excel',    # Modern Excel
    '.ppt': 'powerpoint',
    '.pptx': 'powerpoint',

    # OpenDocument (LibreOffice, etc.)
    '.odt': 'opendocument',  # Writer
    '.ods': 'opendocument',  # Calc
    '.odp': 'opendocument',  # Impress

    # Other document formats
    '.rtf': 'rtf',
    '.epub': 'epub',
}

# -----------------------------------------------------------------------------
# CONSTANTS: Image Formats
IMAGE_EXTENSIONS: set[str] = {
    '.png',   # PNG (lossless, transparency)
    '.jpg',   # JPEG (lossy, photos)
    '.jpeg',  # JPEG (alternative extension)
    '.gif',   # GIF (animations, 256 colors)
    '.bmp',   # Bitmap (uncompressed)
    '.webp',  # WebP (modern, both lossy and lossless)
    '.svg',   # SVG (vector, actually XML)
    '.ico',   # Icons
    '.tiff',  # TIFF (high quality, print)
    '.tif',   # TIFF (alternative extension)
    '.heic',  # HEIC (Apple, high efficiency)
    '.heif',  # HEIF (ISO standard)
}

# =============================================================================
# HELPER FUNCTIONS: File Type Detection
def _detect_file_type(path: Path) -> tuple[str, str]:
    """
    Detects the type of a file based on name and extension.

    This function is the core of the detection system. It checks:
    1. First if the filename is in SPECIAL_FILES (e.g., Makefile, Dockerfile)
    2. Then the extension in CODE_EXTENSIONS
    3. Then in DOCUMENT_EXTENSIONS
    4. Then in IMAGE_EXTENSIONS
    5. Finally assumes plain text as fallback

    Args:
        path: Path object of the file to analyze

    Returns:
        Tuple (category, language/format) where:
        - category: 'code', 'document', 'image', 'text'
        - language: specific identifier (e.g., 'python', 'pdf', 'png')

    Examples:
        >>> _detect_file_type(Path("main.py"))
        ('code', 'python')
        >>> _detect_file_type(Path("Dockerfile"))
        ('code', 'dockerfile')
        >>> _detect_file_type(Path("report.pdf"))
        ('document', 'pdf')
        >>> _detect_file_type(Path("logo.png"))
        ('image', 'png')
    """
    name = path.name
    suffix = path.suffix.lower()

    # 1. Check special files (Makefile, Dockerfile, etc.)
    # These take precedence because the name is more specific than the extension
    if name in SPECIAL_FILES:
        return ('code', SPECIAL_FILES[name])

    # 2. Check source code extensions
    if suffix in CODE_EXTENSIONS:
        return ('code', CODE_EXTENSIONS[suffix])

    # 3. Check documents (PDF, Word, Excel, etc.)
    if suffix in DOCUMENT_EXTENSIONS:
        return ('document', DOCUMENT_EXTENSIONS[suffix])

    # 4. Check images
    if suffix in IMAGE_EXTENSIONS:
        return ('image', suffix[1:])  # Remove the dot (.png -> png)

    # 5. Fallback: assume plain text
    # This allows reading unknown files as text
    return ('text', 'plaintext')

# =============================================================================
# HELPER FUNCTIONS: Text File Reading
def _read_text_file(path: Path) -> str:
    """
    Reads a text file with automatic encoding fallback.

    Problem solved:
    Not all text files are UTF-8. Old files, generated on Windows,
    or with special characters might use other encodings. This function
    tries different encodings in order of probability.

    Attempt order:
    1. utf-8: Modern standard, most common
    2. utf-8-sig: UTF-8 with BOM (Byte Order Mark), common on Windows
    3. latin-1: ISO-8859-1, covers Western Europe
    4. cp1252: Windows-1252, superset of latin-1 used on Windows
    5. iso-8859-1: Alias of latin-1, for safety

    If all fail, reads as binary ignoring errors.

    Args:
        path: Path of the file to read

    Returns:
        File content as string, or error message

    Note:
        latin-1 can never fail because it maps 1:1 with bytes 0-255,
        so in practice we always reach at least that point.
    """
    encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1']

    for encoding in encodings:
        try:
            content = path.read_text(encoding=encoding)
            return content
        except UnicodeDecodeError:
            # This encoding doesn't work, try the next one
            continue
        except Exception as e:
            # Unexpected error (permissions, file disappeared, etc.)
            return f"❌ Read error: {e}"

    # Last attempt: read as binary and decode ignoring errors
    # This may produce � characters but at least reads something
    try:
        content = path.read_bytes().decode('utf-8', errors='replace')
        return content
    except Exception as e:
        return f"❌ Unable to read file: {e}"

def _read_code_file(path: Path, language: str, include_metadata: bool = True) -> str:
    """
    Reads a source code file with optional metadata.

    Unlike simple text reading, this function:
    1. Reads content with encoding fallback
    2. Adds a header with useful metadata for the agent
    3. Formats as markdown code block for syntax highlighting

    The output format is designed to be processed by LLMs:
    - Header with file information
    - Content in code block with specified language
    - Statistics (lines, characters) to evaluate size

    Args:
        path: Path of the file to read
        language: Language identifier (e.g., 'python', 'javascript')
        include_metadata: If True, adds header with metadata

    Returns:
        Formatted content with metadata, or error message

    Example output:
        ```python
        # File: main.py
        # Language: python
        # Lines: 150
        # Size: 4523 chars

        def main():
            ...
        ```
    """
    content = _read_text_file(path)

    # If reading failed, return the error
    if content.startswith("❌"):
        return content

    # If we don't want metadata, return only the content
    if not include_metadata:
        return content

    # Calculate statistics
    lines = content.split('\n')
    line_count = len(lines)
    char_count = len(content)

    # Build formatted output
    # We use triple backticks with language for syntax highlighting
    formatted = f"""```{language}
# File: {path.name}
# Language: {language}
# Lines: {line_count}
# Size: {char_count} chars

{content}
```"""

    return formatted

# =============================================================================
# HELPER FUNCTIONS: Special Document Reading
def _read_pdf(path: Path) -> str:
    """
    Extracts text from a PDF file.

    Uses the pypdf library to extract text from each page.
    Note: pypdf extracts only text, not images or complex layouts.
    For scanned PDFs (images), the extracted text will be empty.

    Dependencies:
        - pypdf: pip install pypdf

    Args:
        path: Path of the PDF file

    Returns:
        Extracted text with page separators, or error message

    Limitations:
        - Does not extract text from images (requires OCR)
        - Complex layouts may lose formatting
        - Tables may not be correctly aligned
    """
    try:
        # Lazy import to not require pypdf if not used
        import pypdf
    except ImportError:
        return "❌ pypdf library not installed. Install with: pip install pypdf"

    try:
        reader = pypdf.PdfReader(str(path))
        pages = []

        for i, page in enumerate(reader.pages, 1):
            # extract_text() may return None for empty pages/images
            text = page.extract_text() or ""
            if text.strip():
                pages.append(f"--- Page {i} ---\n{text}")

        if not pages:
            return (
                "❌ PDF without extractable text.\n"
                "   Possible causes:\n"
                "   - Scanned PDF (contains only images)\n"
                "   - Password protected PDF\n"
                "   - PDF with only graphics/diagrams\n"
                "   Suggestion: use OCR for scanned PDFs"
            )

        # Header with document info
        header = f"📄 PDF: {path.name}\n   Pages: {len(reader.pages)}\n\n"
        return header + "\n\n".join(pages)

    except Exception as e:
        return f"❌ PDF read error: {e}"

def _read_docx(path: Path) -> str:
    """
    Extracts text from a Word document (.docx).

    Uses python-docx to extract:
    - Text paragraphs
    - Table content

    Does not extract:
    - Images
    - Comments/revisions
    - Headers/footers
    - Footnotes

    Dependencies:
        - python-docx: pip install python-docx

    Args:
        path: Path of the .docx file

    Returns:
        Extracted text with tables, or error message

    Note:
        Supports only .docx (Office 2007+), not legacy .doc.
        For .doc, antiword or LibreOffice is needed.
    """
    try:
        from docx import Document
    except ImportError:
        return "❌ python-docx library not installed. Install with: pip install python-docx"

    try:
        doc = Document(str(path))

        # Extract paragraphs (filter empty ones)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Extract tables
        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                # Extract cells and join with |
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                tables.append("\n".join(rows))

        # Build output
        content = "\n\n".join(paragraphs)
        if tables:
            content += "\n\n--- Tables ---\n\n" + "\n\n---\n\n".join(tables)

        # Header with info
        header = (
            f"📄 DOCX: {path.name}\n"
            f"   Paragraphs: {len(paragraphs)}\n"
            f"   Tables: {len(tables)}\n\n"
        )

        return header + content

    except Exception as e:
        return f"❌ DOCX read error: {e}"

def _read_excel(path: Path) -> str:
    """
    Extracts data from an Excel file (.xlsx, .xls).

    Uses pandas to read all sheets from the workbook
    and formats them as text.

    For large files, shows only the first/last 25 rows
    to avoid output that is too long.

    Dependencies:
        - pandas: pip install pandas
        - openpyxl: pip install openpyxl (for .xlsx)
        - xlrd: pip install xlrd (for legacy .xls)

    Args:
        path: Path of the Excel file

    Returns:
        Data formatted as text, or error message
    """
    try:
        import pandas as pd
    except ImportError:
        return "❌ pandas library not installed. Install with: pip install pandas openpyxl"

    try:
        # sheet_name=None reads all sheets as dict
        sheets = pd.read_excel(str(path), sheet_name=None)

        output = [f"📊 Excel: {path.name}", f"   Sheets: {len(sheets)}\n"]

        for name, df in sheets.items():
            output.append(f"\n=== {name} ({len(df)} rows, {len(df.columns)} columns) ===\n")

            # For large sheets, show only head and tail
            if len(df) > 50:
                output.append(df.head(25).to_string())
                output.append(f"\n... {len(df) - 50} rows omitted ...\n")
                output.append(df.tail(25).to_string())
            else:
                output.append(df.to_string())

        return "\n".join(output)

    except Exception as e:
        return f"❌ Excel read error: {e}"

def _read_image(path: Path) -> str:
    """
    Reads information from an image and optionally performs OCR.

    Features:
    1. Extracts base metadata (format, dimensions, color mode)
    2. If pytesseract is available, performs OCR to extract text

    Dependencies:
        - Pillow: pip install Pillow (required)
        - pytesseract: pip install pytesseract (optional, for OCR)
        - tesseract: system installation (brew install tesseract)

    Args:
        path: Path of the image file

    Returns:
        Image information and any OCR text

    Note:
        For SVG, this function returns only basic info.
        SVG is XML and can be read as text with _read_text_file.
    """
    try:
        from PIL import Image
    except ImportError:
        return "❌ Pillow library not installed. Install with: pip install Pillow"

    try:
        # For SVG, read as text (it's XML)
        if path.suffix.lower() == '.svg':
            content = _read_text_file(path)
            return f"🖼️ SVG: {path.name}\n   (XML vector file)\n\n{content}"

        with Image.open(path) as img:
            # Base info
            info = [
                f"🖼️ Image: {path.name}",
                f"   Format: {img.format}",
                f"   Dimensions: {img.size[0]}x{img.size[1]} px",
                f"   Color mode: {img.mode}",
            ]

            # EXIF info if available (photos from cameras)
            if hasattr(img, '_getexif'):
                exif = img._getexif()
                if exif:
                    info.append("   EXIF: present (photo from camera)")

            # Try OCR if available
            try:
                import pytesseract
                text = pytesseract.image_to_string(img)
                if text.strip():
                    info.append(f"\n--- Extracted text (OCR) ---\n{text.strip()}")
                else:
                    info.append("\n   (No text detected in the image)")
            except ImportError:
                info.append("\n   (OCR not available - pip install pytesseract)")
            except Exception as ocr_err:
                info.append(f"\n   (OCR error: {ocr_err})")

        return "\n".join(info)

    except Exception as e:
        return f"❌ Image read error: {e}"

# =============================================================================
# MAIN FUNCTION: File Reading Implementation
def _read_file_impl(
    file_path: str,
    base_dir: Optional[str] = None,
    max_size_mb: float = 50.0,
    include_metadata: bool = True,
    verbose: bool = True
) -> str:
    """
    Core implementation for reading any type of file.

    This is the main function that orchestrates file reading.
    It handles:
    - Security validation (path traversal, sandboxing)
    - File type detection
    - Dispatch to appropriate reader
    - Size limits

    Design pattern:
    - Separation of concerns: this function doesn't read directly,
      it delegates to type-specific _read_* functions
    - Fail-safe: always returns a string, never exceptions
    - Verbose mode: optional colored output for debugging

    Args:
        file_path: File path (absolute or relative)
        base_dir: If specified, the file must be inside this directory
                  (sandboxing for security)
        max_size_mb: Maximum allowed size in MB (default 50MB)
        include_metadata: If True, adds header with file info
        verbose: If True, prints colored output to terminal

    Returns:
        Formatted file content, or error message starting with ❌

    Security:
        - Blocks path traversal (.. in path)
        - Optional sandboxing with base_dir
        - Size limit to prevent memory exhaustion

    Example:
        >>> content = _read_file_impl("src/main.py")
        >>> content = _read_file_impl("data.xlsx", base_dir="./data")
        >>> content = _read_file_impl("huge.log", max_size_mb=100)
    """
    # -------------------------------------------------------------------------
    # STEP 1: Resolve path to absolute path
    path = Path(file_path).resolve()

    # -------------------------------------------------------------------------
    # STEP 2: Security checks
    # 2a. Block path traversal using the resolved path
    # Attack: "../../../etc/passwd" or symlinks to escape the allowed directory
    try:
        # Block path traversal: check for '..' in the original path parts
        # AND verify the resolved path hasn't escaped via symlinks
        if ".." in Path(file_path).parts:
            return "❌ Security error: path traversal (..) not allowed"
    except (ValueError, TypeError):
        return "❌ Security error: invalid file path"

    # Additional check: if base_dir is provided, verify resolved path is within it
    # (this catches symlink-based escapes)
    if base_dir:
        import os
        real_path = os.path.realpath(file_path)
        real_base = os.path.realpath(base_dir)
        if not real_path.startswith(real_base + os.sep) and real_path != real_base:
            return f"❌ Security error: resolved path escapes allowed directory"

    # 2b. Sandboxing: verify the resolved file is inside base_dir
    if base_dir:
        base = Path(base_dir).resolve()
        # Use Path.is_relative_to() for robust containment check
        # (avoids prefix string matching bugs like /tmp/safe vs /tmp/safety)
        try:
            if not path.is_relative_to(base):
                return f"❌ Access denied: file is outside the allowed directory ({base_dir})"
        except (TypeError, ValueError):
            return f"❌ Access denied: file is outside the allowed directory ({base_dir})"

    # -------------------------------------------------------------------------
    # STEP 3: Verify existence and type
    if not path.exists():
        return f"❌ File not found: {path}"

    if not path.is_file():
        if path.is_dir():
            return f"❌ This is a directory, not a file: {path}\n   Use list_directory() to explore directories"
        return f"❌ Not a regular file: {path}"

    # -------------------------------------------------------------------------
    # STEP 4: Verify size
    size_bytes = path.stat().st_size
    size_mb = size_bytes / (1024 * 1024)

    if size_mb > max_size_mb:
        return (
            f"❌ File too large: {size_mb:.1f}MB\n"
            f"   Current limit: {max_size_mb}MB\n"
            f"   Suggestion: increase max_size_mb or read portions of the file"
        )

    # -------------------------------------------------------------------------
    # STEP 5: Detect file type
    category, file_type = _detect_file_type(path)

    if verbose:
        print(f"{Fore.CYAN}📖 Reading: {path.name}")
        print(f"   Type: {category}/{file_type}")
        print(f"   Size: {size_mb:.2f}MB{Style.RESET_ALL}")

    # -------------------------------------------------------------------------
    # STEP 6: Dispatch to appropriate reader
    if category == 'code':
        return _read_code_file(path, file_type, include_metadata)

    elif category == 'document':
        if file_type == 'pdf':
            return _read_pdf(path)
        elif file_type == 'word':
            return _read_docx(path)
        elif file_type in ('excel', 'opendocument'):
            return _read_excel(path)
        else:
            # Other unsupported document formats, try as text
            return _read_text_file(path)

    elif category == 'image':
        return _read_image(path)

    else:
        # Default: read as plain text
        return _read_text_file(path)

# =============================================================================
# FUNCTION: Directory Listing Implementation
def _list_directory_impl(
    dir_path: str,
    pattern: str = "*",
    recursive: bool = False,
    base_dir: Optional[str] = None,
    show_hidden: bool = False,
    max_items: int = 200,
    verbose: bool = True
) -> str:
    """
    Lists the contents of a directory with filters and formatting.

    This function allows agents to explore the file system
    in a secure and structured way. The output is formatted to be
    easily parsable by both humans and LLMs.

    Features:
    - Glob patterns for filtering files (e.g., "*.py", "*.{js,ts}")
    - Optional recursive search
    - Hides hidden files (.) by default
    - Shows file type and size
    - Item limit to avoid output that is too long

    Args:
        dir_path: Directory to explore
        pattern: Glob pattern for filtering (default "*" = all)
                 Supports extended syntax: "*.{py,js}" for multiple extensions
        recursive: If True, also searches in subdirectories
        base_dir: Sandboxing - limits access to this directory
        show_hidden: If True, shows files starting with .
        max_items: Maximum number of items to show
        verbose: If True, colored output to terminal

    Returns:
        Formatted list of files, or error message

    Example output:
        📁 /path/to/project
           Pattern: *.py

          📁 tests/
          📄 main.py  [python] (1.2KB)
          📄 utils.py  [python] (3.4KB)

        📊 Total: 1 folders, 2 files
    """
    # -------------------------------------------------------------------------
    # STEP 1: Resolve and validate path
    path = Path(dir_path).resolve()

    # Security checks (same as _read_file_impl)
    try:
        if ".." in Path(dir_path).parts:
            return "❌ Security error: path traversal (..) not allowed"
    except (ValueError, TypeError):
        return "❌ Security error: invalid directory path"

    if base_dir:
        base = Path(base_dir).resolve()
        try:
            if not path.is_relative_to(base):
                return f"❌ Access denied: directory outside the allowed area ({base_dir})"
        except (TypeError, ValueError):
            return f"❌ Access denied: directory outside the allowed area ({base_dir})"

    if not path.exists():
        return f"❌ Directory not found: {path}"

    if not path.is_dir():
        return f"❌ Not a directory: {path}\n   Use read_file() to read files"

    # -------------------------------------------------------------------------
    # STEP 2: Search files with pattern
    try:
        # Use iterators with a hard limit to avoid materializing huge directories
        MAX_SCAN = max_items * 4  # Scan up to 4x the display limit
        if recursive:
            items = []
            for i, item in enumerate(path.rglob(pattern)):
                if i >= MAX_SCAN:
                    break
                items.append(item)
        else:
            items = []
            for i, item in enumerate(path.glob(pattern)):
                if i >= MAX_SCAN:
                    break
                items.append(item)
    except Exception as e:
        return f"❌ Search error: {e}"

    # -------------------------------------------------------------------------
    # STEP 3: Filter hidden files if requested
    if not show_hidden:
        # Filter items whose name starts with .
        # or that are inside hidden directories
        items = [
            item for item in items
            if not item.name.startswith('.') and
            not any(part.startswith('.') for part in item.parts)
        ]

    # -------------------------------------------------------------------------
    # STEP 4: Separate directories and files, sort alphabetically
    dirs = sorted([item for item in items if item.is_dir()])
    files = sorted([item for item in items if item.is_file()])

    # -------------------------------------------------------------------------
    # STEP 5: Build formatted output
    output = [
        f"📁 {path}",
        f"   Pattern: {pattern}",
        f"   Recursive: {'yes' if recursive else 'no'}",
        ""
    ]

    # Directories (show max half of the limit)
    shown_dirs = 0
    for d in dirs[:max_items // 2]:
        try:
            rel = d.relative_to(path)
            output.append(f"  📁 {rel}/")
            shown_dirs += 1
        except ValueError:
            # relative_to can fail in edge cases
            output.append(f"  📁 {d.name}/")
            shown_dirs += 1

    # Files (the other half of the limit)
    shown_files = 0
    for f in files[:max_items // 2]:
        try:
            rel = f.relative_to(path)
        except ValueError:
            rel = f.name

        # Calculate human-readable size
        try:
            size = f.stat().st_size
            if size < 1024:
                size_str = f"{size}B"
            elif size < 1024 * 1024:
                size_str = f"{size/1024:.1f}KB"
            else:
                size_str = f"{size/(1024*1024):.1f}MB"
        except OSError:
            size_str = "?"

        # Detect file type
        category, ftype = _detect_file_type(f)

        output.append(f"  📄 {rel}  [{ftype}] ({size_str})")
        shown_files += 1

    # -------------------------------------------------------------------------
    # STEP 6: Add summary
    total = len(dirs) + len(files)
    if total > max_items:
        output.append(f"\n  ⚠️  Results truncated: showing {shown_dirs + shown_files} of {total}")
        output.append(f"      Increase max_items or use a more specific pattern")

    output.append(f"\n📊 Total: {len(dirs)} folders, {len(files)} files")

    if verbose:
        print(f"{Fore.CYAN}📁 Scanning: {path}")
        print(f"   Found: {len(dirs)} dirs, {len(files)} files{Style.RESET_ALL}")

    return "\n".join(output)

# =============================================================================
# FUNCTION TOOLS: Wrappers for Agents
@function_tool
def read_file(file_path: str) -> str:
    """
    Reads a file of any supported format.

    This tool allows reading:
    - Source code files (Python, JavaScript, TypeScript, Java, C/C++,
      Rust, Go, Ruby, PHP, Swift, Kotlin, and many others)
    - Documents (PDF, Word .docx, Excel .xlsx)
    - Configuration files (JSON, YAML, TOML, XML, INI)
    - Text and markdown files
    - Images (returns metadata and OCR text if available)

    The tool automatically detects the file type from the extension
    and applies the appropriate parser.

    Args:
        file_path: Path of the file to read (absolute or relative)

    Returns:
        File content with metadata, or error message
        starting with ❌

    Examples:
        read_file("src/main.py")           # Python code
        read_file("config/settings.yaml")  # YAML configuration
        read_file("docs/manual.pdf")       # PDF document
        read_file("data/report.xlsx")      # Excel spreadsheet

    Security:
        - Blocks path traversal (..)
        - 50MB size limit
    """
    # Print visual header for the tool
    print(f"\n{Fore.CYAN}{'='*80}")
    print(f"{Fore.CYAN}🔧 TOOL: read_file")
    print(f"{Fore.CYAN}{'='*80}")
    print(f"{Fore.CYAN}📥 INPUT: {Style.BRIGHT}{file_path}{Style.RESET_ALL}\n")

    result = _read_file_impl(file_path, verbose=True)

    # Print result
    if result.startswith("❌"):
        print(f"{Fore.RED}{result}{Style.RESET_ALL}\n")
    else:
        lines = result.count('\n') + 1
        chars = len(result)
        print(f"{Fore.GREEN}✅ File read: {lines} lines, {chars} characters{Style.RESET_ALL}\n")

    return result

@function_tool
def list_directory(
    directory: str,
    pattern: str = "*",
    recursive: bool = False
) -> str:
    """
    Lists files and folders in a directory.

    Allows exploring the file system with glob filters
    and optional recursive search.

    Args:
        directory: Path of the directory to explore
        pattern: Glob pattern to filter results
                 Examples:
                 - "*" = all files (default)
                 - "*.py" = only Python files
                 - "*.{js,ts}" = JavaScript and TypeScript files
                 - "test_*" = files starting with "test_"
        recursive: If True, also searches in subdirectories

    Returns:
        Formatted list of files with type and size,
        or error message starting with ❌

    Examples:
        list_directory("src")                      # All files in src/
        list_directory("src", "*.py")              # Only Python
        list_directory(".", "*.{py,js}", True)    # Python+JS recursive
    """
    print(f"\n{Fore.CYAN}{'='*80}")
    print(f"{Fore.CYAN}🔧 TOOL: list_directory")
    print(f"{Fore.CYAN}{'='*80}")
    print(f"{Fore.CYAN}📥 INPUT:")
    print(f"{Fore.CYAN}   directory: {Style.BRIGHT}{directory}{Style.NORMAL}")
    print(f"{Fore.CYAN}   pattern: {pattern}")
    print(f"{Fore.CYAN}   recursive: {recursive}{Style.RESET_ALL}\n")

    result = _list_directory_impl(directory, pattern, recursive, verbose=True)

    if result.startswith("❌"):
        print(f"{Fore.RED}{result}{Style.RESET_ALL}\n")
    else:
        print(f"{Fore.GREEN}✅ Directory listed{Style.RESET_ALL}\n")

    return result

@function_tool
def find_code_files(directory: str, language: str) -> str:
    """
    Finds all files of a specific programming language.

    This is a convenience tool that maps the language name
    to the appropriate extensions and searches recursively.

    Supported languages:
    - python, javascript, typescript, java, kotlin
    - rust, go, c, cpp, csharp
    - ruby, php, swift, scala
    - html, css, sql

    Args:
        directory: Directory to search in
        language: Language name (case-insensitive)

    Returns:
        List of found files with path and size

    Examples:
        find_code_files("src", "python")     # All .py files
        find_code_files(".", "typescript")   # All .ts and .tsx files
        find_code_files("app", "rust")       # All .rs files
    """
    # Map language -> glob pattern
    # We use patterns with {} for multiple extensions
    lang_patterns = {
        'python': '*.py',
        'javascript': '*.js',
        'typescript': '*.{ts,tsx}',
        'java': '*.java',
        'kotlin': '*.{kt,kts}',
        'rust': '*.rs',
        'go': '*.go',
        'c': '*.{c,h}',
        'cpp': '*.{cpp,cc,cxx,hpp,hxx}',
        'csharp': '*.cs',
        'ruby': '*.rb',
        'php': '*.php',
        'swift': '*.swift',
        'scala': '*.scala',
        'sql': '*.sql',
        'html': '*.{html,htm}',
        'css': '*.{css,scss,sass,less}',
        'shell': '*.{sh,bash,zsh}',
        'yaml': '*.{yaml,yml}',
        'json': '*.json',
        'markdown': '*.{md,markdown}',
    }

    # Normalize and search the pattern
    lang_lower = language.lower()
    pattern = lang_patterns.get(lang_lower, f'*.{lang_lower}')

    print(f"\n{Fore.CYAN}{'='*80}")
    print(f"{Fore.CYAN}🔧 TOOL: find_code_files")
    print(f"{Fore.CYAN}{'='*80}")
    print(f"{Fore.CYAN}📥 INPUT:")
    print(f"{Fore.CYAN}   directory: {Style.BRIGHT}{directory}{Style.NORMAL}")
    print(f"{Fore.CYAN}   language: {language}")
    print(f"{Fore.CYAN}   resolved pattern: {pattern}{Style.RESET_ALL}\n")

    # Use list_directory with recursive search
    return _list_directory_impl(
        directory,
        pattern=pattern,
        recursive=True,
        verbose=True
    )
